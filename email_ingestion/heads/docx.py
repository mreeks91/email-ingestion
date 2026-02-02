"""DOCX processing head."""

from __future__ import annotations

import io

from email_ingestion.heads.base import HeadInput, HeadResult, Artifact


class DocxHead:
    name = "docx"
    supported_extensions = {"docx"}

    def process(self, head_input: HeadInput) -> HeadResult:
        if not head_input.attachment_bytes:
            return HeadResult()
        try:
            from docx import Document  # type: ignore
        except Exception as exc:  # pragma: no cover - import guard
            raise RuntimeError("Missing dependency: python-docx") from exc
        doc = Document(io.BytesIO(head_input.attachment_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text]
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                table_text.append("\t".join(cells))
        text = "\n".join(paragraphs + table_text).strip() or None
        artifacts = [Artifact(artifact_type="text", text=text)]
        return HeadResult(artifacts=artifacts, metrics={"paragraphs": len(paragraphs)})
